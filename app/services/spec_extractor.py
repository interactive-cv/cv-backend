"""Извлечение текста из файлов ТЗ (PDF, DOCX).

Унифицированный интерфейс: bytes + filename → текст.
Поддерживает таблицы и помечает позиции картинок в DOCX.
"""
import io
from typing import Literal

FileType = Literal["pdf", "docx"]


def detect_file_type(filename: str) -> FileType | None:
    """Определяет тип файла по расширению."""
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        return "pdf"
    if name.endswith(".docx"):
        return "docx"
    return None


def extract_pdf(content: bytes) -> tuple[str, int]:
    """Извлекает текст из PDF. Возвращает (text, pages)."""
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(content))
    pages_text = []
    for page in reader.pages:
        text = page.extract_text() or ""
        pages_text.append(text.strip())
    return "\n\n".join(t for t in pages_text if t), len(reader.pages)


def extract_docx(content: bytes) -> tuple[str, int]:
    """Извлекает текст из DOCX. Возвращает (text, elements_count).

    Извлекает: параграфы (с заголовками), таблицы (как markdown-таблицы),
    помечает позиции картинок [Рисунок N].
    """
    from docx import Document

    doc = Document(io.BytesIO(content))

    # Собираем элементы в порядке следования (параграфы и таблицы перемежаются)
    body = doc.element.body
    elements = []

    # Маппинг XML-элементов → объекты python-docx
    para_idx = 0
    table_idx = 0
    paragraphs = doc.paragraphs
    tables = doc.tables
    image_counter = 0
    element_count = 0

    for child in body.iterchildren():
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        if tag == "p":
            if para_idx < len(paragraphs):
                p = paragraphs[para_idx]
                para_idx += 1
                element_count += 1

                # Проверяем, есть ли картинка в этом параграфе
                has_image = _paragraph_has_image(p)
                if has_image:
                    image_counter += 1

                text = p.text.strip()
                if not text and has_image:
                    elements.append(f"\n[Рисунок {image_counter}]\n")
                elif text and has_image:
                    style_prefix = _style_prefix(p.style.name if p.style else "")
                    elements.append(f"{style_prefix}{text}")
                    elements.append(f"[Рисунок {image_counter}]")
                elif text:
                    style_prefix = _style_prefix(p.style.name if p.style else "")
                    elements.append(f"{style_prefix}{text}")
        elif tag == "tbl":
            if table_idx < len(tables):
                table = tables[table_idx]
                table_idx += 1
                element_count += 1
                md = _table_to_markdown(table)
                if md:
                    elements.append(md)

    text = "\n".join(elements)
    return text, element_count


def _paragraph_has_image(paragraph) -> bool:
    """Проверяет, содержит ли параграф встроенную картинку."""
    # Ищем <a:blip> в run properties (встроенные картинки)
    for run in paragraph.runs:
        if run._element.findall(
            ".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
        ):
            return True
    return False


def _style_prefix(style_name: str) -> str:
    """Превращает имя стиля заголовка в markdown-префикс."""
    name = (style_name or "").lower()
    if "heading 1" in name or name == "heading 1":
        return "## "
    if "heading 2" in name:
        return "### "
    if "heading 3" in name:
        return "#### "
    if "heading" in name:
        return "### "
    if "title" in name:
        return "# "
    return ""


def _table_to_markdown(table) -> str:
    """Превращает таблицу DOCX в markdown-таблицу."""
    rows = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        rows.append("| " + " | ".join(cells) + " |")

    if not rows:
        return ""

    # Добавляем разделитель после первой строки (header)
    if len(rows) >= 1:
        col_count = len(table.rows[0].cells)
        separator = "| " + " | ".join(["---"] * col_count) + " |"
        rows.insert(1, separator)

    return "\n".join(rows)


def extract_spec(filename: str, content: bytes) -> tuple[str, int, str]:
    """Извлекает текст из файла ТЗ.

    Возвращает (text, elements, file_type).
    Выбрасывает ValueError если формат не поддерживается или текст пустой.
    """
    file_type = detect_file_type(filename)
    if not file_type:
        raise ValueError(
            f"Неподдерживаемый формат: {filename}. Поддерживаются PDF и DOCX."
        )

    if not content:
        raise ValueError(f"Пустой файл: {filename}")

    if file_type == "pdf":
        text, count = extract_pdf(content)
    else:
        text, count = extract_docx(content)

    if not text.strip():
        raise ValueError(
            f"Файл {filename} не содержит извлекаемого текста "
            "(возможно, сканы без OCR или только картинки)"
        )

    return text, count, file_type
